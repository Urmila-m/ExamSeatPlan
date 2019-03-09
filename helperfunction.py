import pandas as pd
import numpy as np
import math
import docx

def partition(size,num_partion):
    step = size/num_partion
    result = list()
    ans = 0
    for i in range(num_partion):
        ans += step
        result.append(int(ans-1))

    return result

# def save_todb(exam_name, result_df):
#     import sqlalchemy as db
#     from config import database_ip, database_password, database_username
#     from sqlalchemy_utils import database_exists, create_database
#     database_name = exam_name
#     database_engine = db.create_engine('mysql+mysqlconnector://{0}:{1}@{2}/{3}'.
#                                                 format(database_username, database_password, 
#                                                         database_ip, database_name))
#     if not database_exists(database_engine.url):
#         create_database(database_engine.url)
    
#     for fname, planed_hall in result_df:
#         planed_hall.to_sql(con=database_engine, name=fname, if_exists='replace', index = False)

def Grid2List(table):
    table_iter = table.iteritems()
    lst = []
    next(table_iter)
    for column, row in table_iter:
        for row_index, rowdata in enumerate(row):
            if rowdata != 'nan':
                Seat = 'C' + str(column) + 'R' + str(row_index+1)
                lst.append([rowdata, Seat])
    
    list_table = pd.DataFrame(lst, columns = ['Name', 'Seat'])
    list_table.dropna(inplace=True)
    return list_table

# def get_exams():
#     import sqlalchemy as db
#     from config import database_ip, database_password, database_username


#     database_engine = db.create_engine('mysql+mysqlconnector://{0}:{1}@{2}'.
#                                                 format(database_username, database_password, 
#                                                         database_ip))
    
#     for exam_name in database_engine.execute('show databases').fetchall():
#         yield exam_name[0]      


  

def equalise(dataframes, totalseat):
    # Merges dataframes so that all student occupy all the seats
    num_exam = len(dataframes)
    sizes = list(map(len, dataframes))
    total_size = sum(sizes)
    assert(total_size <= totalseat)

    max_size = max(sizes)

    if totalseat/num_exam >= max_size:
        group_size = max_size
    
    else:
        group_size = math.ceil(totalseat/num_exam) #up ceil

    total_dataframes = pd.DataFrame(np.array(["nan"]* (group_size *num_exam)), columns = ['Key']) # initializing

    # make it more readable
    a = 0
    for i, df in enumerate(dataframes):
        b = a + sizes[i]
        total_dataframes.Key[a:b] = df.Key[0:sizes[i]]
        a = b
    
    ans_dataframe = list()
    a = 0
    for i in range(num_exam):
        b = (i+1) * group_size
        ans_dataframe.append(total_dataframes[a:b])
        a = b
    
    return ans_dataframe

def arrange_seat(equalised_dataframes):
    group_num = len(equalised_dataframes)
    group_size = len(equalised_dataframes[0])
    
    total_list = list()
    for df in equalised_dataframes:
        total_list.append(df.values)
    
    total_list = np.array(total_list)
    dataframe = pd.DataFrame(total_list.T.reshape(group_size, group_num))
    dataframe = dataframe.values.reshape(-1,1)
    dataframe = pd.DataFrame(dataframe,columns = ['Key'])
    dataframe['Key'] = dataframe['Key'].astype(str)
    return dataframe

def plan_examhall(dataframe, row, column):    
    dataframe['Row'] = dataframe.index/(len(dataframe)/(row)) + 1
    dataframe['Row'] = dataframe['Row'].apply(lambda x: int(x))
     
    dataframe['Column'] = dataframe.index % column + 1
    dataframe = dataframe.set_index(['Column','Row'])
    return pd.pivot_table(dataframe, values='Key', index=['Row'],
                    columns=['Column'], aggfunc=lambda x: ' '.join(x))

def result_2docx(df, examination_name ,examhall, date, time, location):

    doc = docx.Document()
    exam_name = doc.add_paragraph()
    exam_name.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
    exam_name.add_run(examination_name).bold = True

    date_para = doc.add_paragraph()
    date_para.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.RIGHT
    date_para.add_run(date)

    exam_hall_para = doc.add_paragraph()
    exam_hall_para.add_run('Examination Hall:- ')
    exam_hall_para.add_run(examhall)

    time_para = doc.add_paragraph()
    time_para.add_run('Time:- ')
    time_para.add_run(time)
    # add a table to the end and create a reference variable
    # extra row is so we can add the header row
    t = doc.add_table(df.shape[0]+1, df.shape[1])

    # add the header rows.
    for j in range(df.shape[1]):
        t.cell(0,j).text = str(df.columns[j])
        

    # add the rest of the data frame
    for i in range(df.shape[0]):
        for j in range(df.shape[1]):
            t.cell(i+1,j).text = str(df.values[i,j])

    # save the doc
    doc.save(location)
    
